from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# Colour palette
NAVY  = RGBColor(0x1F, 0x35, 0x64)
BLUE  = RGBColor(0x27, 0x63, 0xAB)
TEAL  = RGBColor(0x17, 0x7B, 0x7B)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)


def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    pPr.append(shd)


def shade_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)


def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement('w:' + edge)
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '6')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '2763AB')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, color=DARK, size=11):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.name   = 'Calibri'
    r.font.size   = Pt(size)
    r.font.color.rgb = color
    return r


def add_h2(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'E8F0FB')
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(text)
    r.bold = True
    r.font.name  = 'Calibri'
    r.font.size  = Pt(16)
    r.font.color.rgb = BLUE
    return p


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name  = 'Calibri'
    r.font.size  = Pt(13)
    r.font.color.rgb = TEAL
    return p


def add_body():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    return p


def new_bullet(parts):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name  = 'Calibri'
        r.font.size  = Pt(11)
        r.font.color.rgb = DARK
    return p


def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2763AB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)


STATUS_COLOR = {
    'SOLVED':      RGBColor(0x1A, 0x73, 0x48),
    'IN PROGRESS': RGBColor(0xB4, 0x5F, 0x06),
    'UNDERSTOOD':  RGBColor(0x1A, 0x73, 0x48),
}
STATUS_LABEL = {
    'SOLVED':      'Status:  SOLVED',
    'IN PROGRESS': 'Status:  IN PROGRESS - Testing WebScraping.ai residential Israeli IP',
    'UNDERSTOOD':  'Status:  UNDERSTOOD - Must render full page',
}


def status_line(key):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(STATUS_LABEL[key])
    r.bold = True
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11)
    r.font.color.rgb = STATUS_COLOR[key]


# ──────────────────────────────────────────────────────────────────────────────
# TITLE BLOCK
# ──────────────────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(title_para, '1F3564')
r = title_para.add_run('LoFrayer \u2014 Scraping Obstacles Log')
r.bold = True
r.font.name  = 'Calibri'
r.font.size  = Pt(26)
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
title_para.paragraph_format.space_before = Pt(14)
title_para.paragraph_format.space_after  = Pt(14)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(subtitle_para, '2763AB')
rs = subtitle_para.add_run('Technical Log  \u2022  Scraping Infrastructure  \u2022  March 2026')
rs.font.name  = 'Calibri'
rs.font.size  = Pt(12)
rs.font.color.rgb = RGBColor(0xD0, 0xE4, 0xFF)
subtitle_para.paragraph_format.space_before = Pt(6)
subtitle_para.paragraph_format.space_after  = Pt(18)

# ──────────────────────────────────────────────────────────────────────────────
# PROJECT OVERVIEW
# ──────────────────────────────────────────────────────────────────────────────
add_h2('Project Overview')
p = add_body()
add_run(p, 'LoFrayer', bold=True)
add_run(p, ' is an Israeli app centralising membership discounts. The main scraping challenge is getting discount data from ')
add_run(p, 'benefits.isracard.co.il', bold=True)
add_run(p, ' \u2014 a React SPA that is both ')
add_run(p, 'geo-blocked (Israel only)', bold=True)
add_run(p, ' and protected by ')
add_run(p, 'Cloudflare bot detection', bold=True)
add_run(p, '.')
add_hr()

# ──────────────────────────────────────────────────────────────────────────────
# SESSION 1
# ──────────────────────────────────────────────────────────────────────────────
add_h2('Session 1 \u2014 Obstacles & Solutions')

# Obstacle 1
add_h3('Obstacle 1: Geo-Blocking')
new_bullet([('Error: ', True),  ('Page not loading from non-Israeli IP', False)])
new_bullet([('Cause: ', True),  ('benefits.isracard.co.il blocks all non-Israeli IP addresses', False)])
new_bullet([('Solution: ', True),
            ('Set up Google Cloud VM in zone ', False),
            ('me-west1-a', True),
            (' (Israel region) with a real Israeli IP address (', False),
            ('34.165.116.90', True),
            ('), running Browserless Chrome Docker container', False)])
status_line('SOLVED')

# Obstacle 2
add_h3('Obstacle 2: Browserless Docker Container Stopped')
new_bullet([('Error: ', True),
            ('"Isracard scraping failed \u2014 check Browserless Docker status on VM"', False)])
new_bullet([('Cause: ', True), ('The Docker container had stopped running on the VM', False)])
new_bullet([('Solution: ', True),
            ('SSH into VM and run ', False),
            ('sudo docker start browserless', True)])
status_line('SOLVED')

# Obstacle 3
add_h3('Obstacle 3: Browserless /function Endpoint Broken')
new_bullet([('Error: ', True), ('"Unexpected end of JSON input"', False)])
new_bullet([('Cause: ', True),
            ('This version of Browserless (Puppeteer v21.9.0 / Chrome 121) returns an empty body on the ', False),
            ('/function', True),
            (' endpoint', False)])
new_bullet([('Solution: ', True),
            ('Switched to the ', False),
            ('/content', True),
            (' endpoint which works correctly and returns rendered HTML', False)])
status_line('SOLVED')

# Obstacle 4
add_h3('Obstacle 4: Cloudflare Bot Detection (Core Problem)')
new_bullet([('Error: ', True),
            ('HTML response was only 5,562 chars \u2014 "Attention Required! | Cloudflare"', False)])
new_bullet([('Cause: ', True),
            ('Even with an Israeli IP, Cloudflare detects ', False),
            ('datacenter IPs', True),
            (' (Google Cloud) as bots. ', False),
            ('Residential IPs', True),
            (' are required.', False)])
p_fix = doc.add_paragraph(style='List Bullet')
p_fix.paragraph_format.space_before = Pt(2)
p_fix.paragraph_format.space_after  = Pt(2)
add_run(p_fix, 'Attempted Fixes:', bold=True)
fixes = [
    'Added stealth: true to Browserless request \u2192 Failed (see Obstacle 5)',
    'Investigated DevTools Network tab for hidden API \u2192 No API found (see Obstacle 6)',
    'Switching to WebScraping.ai residential proxy \u2192 Currently testing',
]
for fix in fixes:
    p = doc.add_paragraph(style='List Bullet 2')
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(fix)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
status_line('IN PROGRESS')

# Obstacle 5
add_h3('Obstacle 5: stealth:true Not Supported in Browserless Version')
new_bullet([('Error: ', True), ('HTTP 400 \u2014 "\\"stealth\\" is not allowed"', False)])
new_bullet([('Cause: ', True),
            ('The installed version of Browserless (Puppeteer v21.9.0) does not support the ', False),
            ('stealth', True),
            (' parameter', False)])
new_bullet([('Solution: ', True), ('Removed stealth: true from the request body', False)])
status_line('SOLVED')

# Obstacle 6
add_h3('Obstacle 6: No Direct API Endpoint for Isracard')
new_bullet([('Investigation: ', True),
            ('Opened DevTools Network tab and filtered through ', False),
            ('174+ network requests', True),
            (' on benefits.isracard.co.il', False)])
p_f = doc.add_paragraph(style='List Bullet')
p_f.paragraph_format.space_before = Pt(2)
p_f.paragraph_format.space_after  = Pt(2)
add_run(p_f, 'Findings:', bold=True)
finding_data = [
    ('First large request (460361382.json, 11.5 kB) was a QuadPixel tracking SDK \u2014 ',
     'not discount data', ''),
    ('All other requests were analytics/tracking (Google, Taboola, DoubleClick)', '', ''),
    ('The ', 'online-benefits/', ' document (30.6 kB) IS the main page with server-side rendered content'),
]
for f in finding_data:
    p = doc.add_paragraph(style='List Bullet 2')
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(f[0])
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK
    if f[1]:
        rb = p.add_run(f[1])
        rb.bold = True
        rb.font.name = 'Calibri'; rb.font.size = Pt(11); rb.font.color.rgb = DARK
    if f[2]:
        r2 = p.add_run(f[2])
        r2.font.name = 'Calibri'; r2.font.size = Pt(11); r2.font.color.rgb = DARK
p_conc = doc.add_paragraph(style='List Bullet')
p_conc.paragraph_format.space_before = Pt(2)
add_run(p_conc, 'Conclusion: ', bold=True)
add_run(p_conc, 'The discount data is server-side rendered directly into the HTML. There is ')
add_run(p_conc, 'no separate API endpoint', bold=True)
add_run(p_conc, ' to call. We ')
add_run(p_conc, 'MUST load the full page', bold=True)
add_run(p_conc, '.')
status_line('UNDERSTOOD')

# Obstacle 7
add_h3('Obstacle 7: WebScraping.ai Wrong API Key Format')
new_bullet([('Error: ', True), ('403 Wrong API key', False)])
new_bullet([('Cause: ', True),
            ('The API key stored in Supabase was in ', False),
            ('hex format', True),
            (', but the real WebScraping.ai key is in ', False),
            ('UUID format', True)])
new_bullet([('Solution: ', True),
            ('Deleted old ', False),
            ('WEBSCRAPING_AI_KEY', True),
            (' secret in Supabase dashboard, added new one with correct UUID value', False)])
status_line('SOLVED')

add_hr()

# ──────────────────────────────────────────────────────────────────────────────
# ARCHITECTURE TABLE
# ──────────────────────────────────────────────────────────────────────────────
add_h2('Architecture Decisions')
doc.add_paragraph()  # small spacer

table_data = [
    ('Component',       'Choice',                                'Reason'),
    ('Scraping engine', 'WebScraping.ai (residential IL proxy)', 'Bypasses both geo-block and Cloudflare'),
    ('AI parsing',      'Google Gemini 2.0 Flash',               'Free tier, good Hebrew understanding'),
    ('Database',        'Supabase (PostgreSQL)',                  'Already in use for the app'),
    ('Deployment',      'Supabase Edge Functions (Deno)',         'Serverless, no server management'),
    ('VM',              'Google Cloud me-west1-a (Israel)',       'Israeli IP for geo-restricted sites'),
]

tbl = doc.add_table(rows=len(table_data), cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(4.5), Cm(6.5), Cm(6.5)]

for i, row in enumerate(tbl.rows):
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]
        text = table_data[i][j]
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)
        para.paragraph_format.left_indent  = Cm(0.2)
        if i == 0:
            shade_cell(cell, '2763AB')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(text)
            r.bold = True
            r.font.name  = 'Calibri'
            r.font.size  = Pt(11)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        else:
            fill = 'F5F8FF' if i % 2 == 0 else 'FFFFFF'
            shade_cell(cell, fill)
            if j == 0:
                r = para.add_run(text)
                r.bold = True
                r.font.name  = 'Calibri'
                r.font.size  = Pt(11)
                r.font.color.rgb = BLUE
            else:
                r = para.add_run(text)
                r.font.name  = 'Calibri'
                r.font.size  = Pt(11)
                r.font.color.rgb = DARK
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

add_hr()

# ──────────────────────────────────────────────────────────────────────────────
# KEY LESSONS
# ──────────────────────────────────────────────────────────────────────────────
add_h2('Key Lessons Learned')

lessons = [
    ('Residential IPs beat Datacenter IPs',
     ' \u2014 Cloudflare trusts residential IPs far more than cloud datacenter IPs, even if both are in Israel'),
    ('Docker containers need monitoring',
     ' \u2014 They can stop unexpectedly; need health checks'),
    ('Always check DevTools Network tab',
     ' \u2014 Before building complex scraping infrastructure, verify if there\'s a simpler API endpoint'),
    ('API key format matters',
     ' \u2014 Always verify the expected format (hex vs UUID vs other) against the service documentation'),
    ('Server-side rendering requires full page load',
     ' \u2014 If data is baked into HTML (not loaded via API), you must execute JavaScript and wait for the page to render'),
]

for title, detail in lessons:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rb = p.add_run(title)
    rb.bold = True
    rb.font.name  = 'Calibri'
    rb.font.size  = Pt(11)
    rb.font.color.rgb = NAVY
    rd = p.add_run(detail)
    rd.font.name  = 'Calibri'
    rd.font.size  = Pt(11)
    rd.font.color.rgb = DARK

add_hr()

# ──────────────────────────────────────────────────────────────────────────────
# CURRENT STATUS
# ──────────────────────────────────────────────────────────────────────────────
add_h2('Current Status (as of March 2026)')
p = add_body()
add_run(p, 'WebScraping.ai', bold=True)
add_run(p, ' with ')
add_run(p, 'country=il', bold=True)
add_run(p, ', ')
add_run(p, 'render_js=1', bold=True)
add_run(p, ', ')
add_run(p, 'proxy=residential', bold=True)
add_run(p, ' deployed to Supabase Edge Function ')
add_run(p, 'scrape-oracle', bold=True)
add_run(p, '.')

p2 = add_body()
add_run(p2, 'Testing in progress.', italic=True, color=RGBColor(0x55, 0x66, 0x88))

# ──────────────────────────────────────────────────────────────────────────────
# SAVE
# ──────────────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\Win11\OneDrive\Dokumenti\lofrayerobstacels.docx'
doc.save(out_path)
print('Saved:', out_path)
